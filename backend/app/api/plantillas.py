# backend/app/api/plantillas.py
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query
from sqlalchemy.orm import Session
from typing import List, Dict, Any, Optional
import re, io, zipfile
from fastapi.responses import FileResponse

from app.db.session import get_global_db
from app.core.dependencies import get_current_active_user, check_project_access
from app.models.global_models import Usuario, Plantilla, PlantillaCampo
from app.db.router import get_project_db
from app.services.log_service import registrar_log
from pydantic import BaseModel

import tempfile
import subprocess
import shutil
import platform
from datetime import datetime
from pathlib import Path

import os as _os

router = APIRouter()

# ── Schemas ───────────────────────────────────────────────────────────────────

class PlantillaCreate(BaseModel):
    id_proyecto: int
    nombre: str
    descripcion: Optional[str] = None
    origen: str = "editor"

class PlantillaUpdate(BaseModel):
    nombre: Optional[str] = None
    descripcion: Optional[str] = None
    activa: Optional[bool] = None

class CampoMapeo(BaseModel):
    placeholder: str
    campo_bd: str
    orden: int = 0

class MapeoRequest(BaseModel):
    campos: List[CampoMapeo]

class PreviewRequest(BaseModel):
    placeholders: Dict[str, str] = {}  
    highlight_placeholders: bool = True  

class PreviewPDFRequest(BaseModel):
    placeholders: Dict[str, str] = {} 
    preview_on: bool = False

# ── Helpers ───────────────────────────────────────────────────────────────────

def _require_analista(user: Usuario):
    """
    FIX: user.rol es un objeto ORM (Rol), no un string.
    Se accede a user.rol.nombre para comparar.
    """
    rol_nombre = user.rol.nombre if hasattr(user.rol, "nombre") else str(user.rol)
    if rol_nombre not in ("superadmin", "analista"):
        raise HTTPException(status_code=403, detail="Requiere rol analista o superadmin.")

def _get_plantilla_or_404(db: Session, pid: int) -> Plantilla:
    p = db.query(Plantilla).filter(Plantilla.id == pid).first()
    if not p:
        raise HTTPException(status_code=404, detail="Plantilla no encontrada.")
    return p

def _slug_from_proyecto_id(db: Session, proyecto_id: int) -> str:
    from sqlalchemy import text
    row = db.execute(
        text("SELECT slug FROM proyectos WHERE id = :id"), {"id": proyecto_id}
    ).first()
    if not row:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado.")
    return row.slug

def _get_campos_analisis(slug: str) -> List[str]:
    """
    FIX: Busca primero tabla_analisis (que siempre existe si hay padrón cargado).
    Si no existe, intenta tabla_temporal como fallback.
    """
    for tabla in ("tabla_analisis", "tabla_temporal"):
        try:
            db = next(get_project_db(slug))
            from sqlalchemy import text
            rows = db.execute(text(f"SHOW COLUMNS FROM `{tabla}`")).fetchall()
            campos = [r[0] for r in rows if not r[0].startswith("_")]
            if campos:
                return campos
        except Exception:
            continue
    return []

def _extraer_placeholders_docx(contenido: bytes) -> List[str]:
    """
    Extrae placeholders de un .docx:
    1. {{campo}} (formato simple)
    2. MERGEFIELD campo (combinación de correspondencia de Word)
    3. DOCPROPERTY campo (propiedades de documento)
    """
    placeholders = []
    seen = set()
    try:
        with zipfile.ZipFile(io.BytesIO(contenido)) as z:
            with z.open("word/document.xml") as f:
                xml_text = f.read().decode("utf-8", errors="replace")

            texto_plano = re.sub(r"<[^>]+>", " ", xml_text)

            # 1. Buscar {{campo}}
            for m in re.finditer(r"\{\{(\w+)\}\}", texto_plano):
                ph = m.group(1)
                if ph not in seen:
                    seen.add(ph)
                    placeholders.append(ph)

            # 2. Buscar MERGEFIELD
            for m in re.finditer(
                r'MERGEFIELD\s+"?([^"<\s\\]+)"?',
                xml_text,
                re.IGNORECASE
            ):
                ph = m.group(1).strip()
                if ph and ph not in seen:
                    seen.add(ph)
                    placeholders.append(ph)

            # 3. Buscar campos de formulario / propiedades
            for m in re.finditer(
                r'DOCPROPERTY\s+"?([^"<\s\\]+)"?',
                xml_text,
                re.IGNORECASE
            ):
                ph = m.group(1).strip()
                if ph and ph not in seen:
                    seen.add(ph)
                    placeholders.append(ph)

    except Exception as e:
        print(f"[plantillas] Error extrayendo placeholders: {e}")

    return placeholders

def _mapeo_automatico(
    placeholders: List[str], campos_bd: List[str]
) -> Dict[str, Optional[str]]:
    campos_idx = {c.lower(): c for c in campos_bd}
    return {ph: campos_idx.get(ph.lower()) for ph in placeholders}


@router.get("/")
def listar_plantillas(
    proyecto_id: Optional[int] = Query(None),
    activa: Optional[bool] = Query(None),
    current_user: Usuario = Depends(get_current_active_user),
    db: Session = Depends(get_global_db),
):
    from sqlalchemy import text

    if current_user.rol.nombre == "superadmin":
        ids_permitidos = None
    else:
        rows = db.execute(
            text("SELECT id_proyecto FROM usuario_proyecto WHERE id_usuario = :u"),
            {"u": current_user.id},
        ).fetchall()
        ids_permitidos = [r.id_proyecto for r in rows]

    q = db.query(Plantilla)
    if proyecto_id:
        q = q.filter(Plantilla.id_proyecto == proyecto_id)
    if activa is not None:
        q = q.filter(Plantilla.activa == activa)
    if ids_permitidos is not None:
        q = q.filter(Plantilla.id_proyecto.in_(ids_permitidos))

    plantillas = q.order_by(Plantilla.created_at.desc()).all()

    result = []
    for p in plantillas:
        proy = db.execute(
            text("SELECT nombre, slug FROM proyectos WHERE id = :id"), {"id": p.id_proyecto}
        ).first()
        total_campos = (
            db.query(PlantillaCampo)
            .filter(PlantillaCampo.id_plantilla == p.id)
            .count()
        )
        result.append({
            "id": p.id,
            "id_proyecto": p.id_proyecto,
            "proyecto_nombre": proy.nombre if proy else "—",
            "proyecto_slug": proy.slug if proy else "",
            "nombre": p.nombre,
            "descripcion": p.descripcion,
            "origen": p.origen,
            "activa": p.activa,
            "ruta_archivo": p.ruta_archivo,
            "created_at": p.created_at.isoformat() if p.created_at else None,
            "updated_at": p.updated_at.isoformat() if p.updated_at else None,
            "total_campos": total_campos,
        })
    return result

# ── POST / — crear metadata (para editor desde cero) ─────────────────────────

@router.post("/")
def crear_plantilla(
    body: PlantillaCreate,
    current_user: Usuario = Depends(get_current_active_user),
    db: Session = Depends(get_global_db),
):
    _require_analista(current_user)
    p = Plantilla(
        id_proyecto=body.id_proyecto,
        nombre=body.nombre,
        descripcion=body.descripcion,
        origen=body.origen,
        activa=True,
        created_by=current_user.id,
    )
    db.add(p)
    db.commit()
    db.refresh(p)
    registrar_log(
        db, current_user.id, "crear_plantilla",
        f"Plantilla '{p.nombre}' creada (proyecto {body.id_proyecto})",
        body.id_proyecto,
    )
    return {"id": p.id, "mensaje": "Plantilla creada."}

# ── POST /subir — subir .docx ─────────────────────────────────────────────────

@router.post("/subir")
async def subir_plantilla_docx(
    proyecto_id: int = Query(...),
    nombre: str = Query(...),
    descripcion: Optional[str] = Query(None),
    file: UploadFile = File(...),
    current_user: Usuario = Depends(get_current_active_user),
    db: Session = Depends(get_global_db),
):
    _require_analista(current_user)

    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Solo se aceptan archivos .docx.")

    contenido = await file.read()

    placeholders = _extraer_placeholders_docx(contenido)

    upload_dir = _os.path.join("uploads", "plantillas")
    _os.makedirs(upload_dir, exist_ok=True)
    safe_name = re.sub(r"[^\w.\-]", "_", file.filename)
    ruta = _os.path.join(upload_dir, f"{proyecto_id}_{safe_name}")
    with open(ruta, "wb") as f_out:
        f_out.write(contenido)

    p = Plantilla(
        id_proyecto=proyecto_id,
        nombre=nombre,
        descripcion=descripcion,
        origen="upload",
        ruta_archivo=ruta,
        activa=True,
        created_by=current_user.id,
    )
    db.add(p)
    db.commit()
    db.refresh(p)

    slug = _slug_from_proyecto_id(db, proyecto_id)
    campos_bd = _get_campos_analisis(slug)
    mapeo_auto = _mapeo_automatico(placeholders, campos_bd)

    for orden, (ph, campo) in enumerate(mapeo_auto.items()):
        if campo:
            db.add(PlantillaCampo(
                id_plantilla=p.id,
                placeholder=f"{{{{{ph}}}}}",
                campo_bd=campo,
                orden=orden,
            ))
    db.commit()

    registrar_log(
        db, current_user.id, "subir_plantilla",
        f"Plantilla '{nombre}' subida. {len(placeholders)} placeholders detectados.",
        proyecto_id,
    )

    return {
        "id": p.id,
        "mensaje": f"Plantilla subida con {len(placeholders)} placeholders detectados.",
        "placeholders": placeholders,
        "mapeo_automatico": mapeo_auto,
        "campos_disponibles": campos_bd,
    }

# ── GET /campos-temporales-slug/{proyecto_slug} ───────────────────────────────

@router.get("/{proyecto_slug}/campos-temporales-slug")
def campos_temporales_por_slug(
    proyecto_slug: str,
    current_user: Usuario = Depends(get_current_active_user),
    db: Session = Depends(get_global_db),
):
    campos = _get_campos_analisis(proyecto_slug)
    return {"campos": campos, "proyecto_slug": proyecto_slug}

# ── GET /{plantilla_id} — detalle ─────────────────────────────────────────────

@router.get("/{plantilla_id}")
def detalle_plantilla(
    plantilla_id: int,
    current_user: Usuario = Depends(get_current_active_user),
    db: Session = Depends(get_global_db),
):
    p = _get_plantilla_or_404(db, plantilla_id)
    slug = _slug_from_proyecto_id(db, p.id_proyecto)
    campos = (
        db.query(PlantillaCampo)
        .filter(PlantillaCampo.id_plantilla == plantilla_id)
        .order_by(PlantillaCampo.orden)
        .all()
    )
    return {
        "id": p.id,
        "id_proyecto": p.id_proyecto,
        "proyecto_slug": slug,
        "nombre": p.nombre,
        "descripcion": p.descripcion,
        "origen": p.origen,
        "activa": p.activa,
        "ruta_archivo": p.ruta_archivo,
        "created_at": p.created_at.isoformat() if p.created_at else None,
        "campos": [
            {"id": c.id, "placeholder": c.placeholder, "campo_bd": c.campo_bd, "orden": c.orden}
            for c in campos
        ],
    }

# ── PUT /{plantilla_id} ───────────────────────────────────────────────────────

@router.put("/{plantilla_id}")
def actualizar_plantilla(
    plantilla_id: int,
    body: PlantillaUpdate,
    current_user: Usuario = Depends(get_current_active_user),
    db: Session = Depends(get_global_db),
):
    _require_analista(current_user)
    p = _get_plantilla_or_404(db, plantilla_id)
    if body.nombre is not None:
        p.nombre = body.nombre
    if body.descripcion is not None:
        p.descripcion = body.descripcion
    if body.activa is not None:
        p.activa = body.activa
    db.commit()
    return {"mensaje": "Actualizada."}

# ── DELETE /{plantilla_id} ────────────────────────────────────────────────────

@router.delete("/{plantilla_id}")
def eliminar_plantilla(
    plantilla_id: int,
    current_user: Usuario = Depends(get_current_active_user),
    db: Session = Depends(get_global_db),
):
    _require_analista(current_user)
    p = _get_plantilla_or_404(db, plantilla_id)
    # Soft delete: desactivar en lugar de borrar físicamente
    p.activa = False
    db.commit()
    return {"mensaje": "Plantilla desactivada."}

# ── POST /{plantilla_id}/mapear ───────────────────────────────────────────────

@router.post("/{plantilla_id}/mapear")
def guardar_mapeo(
    plantilla_id: int,
    body: MapeoRequest,
    current_user: Usuario = Depends(get_current_active_user),
    db: Session = Depends(get_global_db),
):
    _require_analista(current_user)
    p = _get_plantilla_or_404(db, plantilla_id)
    # Reemplazar mapeo anterior completamente
    db.query(PlantillaCampo).filter(PlantillaCampo.id_plantilla == plantilla_id).delete()
    for c in body.campos:
        db.add(PlantillaCampo(
            id_plantilla=plantilla_id,
            placeholder=c.placeholder,
            campo_bd=c.campo_bd,
            orden=c.orden,
        ))
    db.commit()
    registrar_log(
        db, current_user.id, "guardar_mapeo",
        f"Mapeo actualizado: plantilla {plantilla_id}, {len(body.campos)} campos.",
        p.id_proyecto,
    )
    return {"mensaje": f"{len(body.campos)} campos guardados."}

# ── GET /{plantilla_id}/campos-temporales ─────────────────────────────────────

@router.get("/{plantilla_id}/campos-temporales")
def campos_temporales(
    plantilla_id: int,
    current_user: Usuario = Depends(get_current_active_user),
    db: Session = Depends(get_global_db),
):
    p = _get_plantilla_or_404(db, plantilla_id)
    slug = _slug_from_proyecto_id(db, p.id_proyecto)
    return {"campos": _get_campos_analisis(slug), "proyecto_slug": slug}

# ── GET /{plantilla_id}/preview-mapeo ────────────────────────────────────────

@router.get("/{plantilla_id}/preview-mapeo")
def preview_mapeo(
    plantilla_id: int,
    current_user: Usuario = Depends(get_current_active_user),
    db: Session = Depends(get_global_db),
):
    p = _get_plantilla_or_404(db, plantilla_id)
    slug = _slug_from_proyecto_id(db, p.id_proyecto)
    campos_bd = _get_campos_analisis(slug)
    campos_actuales = (
        db.query(PlantillaCampo)
        .filter(PlantillaCampo.id_plantilla == plantilla_id)
        .order_by(PlantillaCampo.orden)
        .all()
    )
    return {
        "campos_actuales": [
            {"placeholder": c.placeholder, "campo_bd": c.campo_bd}
            for c in campos_actuales
        ],
        "campos_disponibles": campos_bd,
    }

@router.get("/{plantilla_id}/descargar")
def descargar_plantilla(
    plantilla_id: int,
    current_user: Usuario = Depends(get_current_active_user),
    db: Session = Depends(get_global_db),
):
    """Descarga el archivo .docx de una plantilla para edición."""
    p = _get_plantilla_or_404(db, plantilla_id)
    if not p.ruta_archivo or not _os.path.exists(p.ruta_archivo):
        raise HTTPException(status_code=404, detail="Archivo no encontrado en el servidor.")
    return FileResponse(
        p.ruta_archivo,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{p.nombre}.docx"
    )

@router.post("/{plantilla_id}/preview")
def preview_plantilla_html(
    plantilla_id: int,
    body: PreviewRequest,
    current_user: Usuario = Depends(get_current_active_user),
    db: Session = Depends(get_global_db),
):
    """
    Genera vista previa HTML de la plantilla.
    - Modo OFF (placeholders vacío): muestra {{campos}} resaltados para identificarlos.
    - Modo ON (placeholders con datos): reemplaza con valores de ejemplo, 
      marca textos largos como posible desborde.
    """
    p = _get_plantilla_or_404(db, plantilla_id)
    
    if not p.ruta_archivo or not _os.path.exists(p.ruta_archivo):
        raise HTTPException(status_code=404, detail="Archivo de plantilla no encontrado.")
    
    try:
        import docx
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        import html as html_module
        from lxml import etree
        import re
        import base64
        from io import BytesIO
        
        doc = docx.Document(p.ruta_archivo)
        preview_on = bool(body.placeholders)
        replacements = body.placeholders if preview_on else {}
        
        # ── Extraer imagen del XML para obtener posición y behindText ──────
        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        }
        
        # Mapa de relaciones para imágenes
        image_rels = {}
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                image_rels[rel_id] = {
                    'blob': rel.target_part.blob,
                    'mime': rel.target_part.content_type,
                }
        
        def extract_image_info(drawing_xml) -> dict:
            """Extrae posición y tamaño de un elemento <w:drawing>."""
            info = {
                'rId': None,
                'width_emu': 0,
                'height_emu': 0,
                'behind_text': False,
                'floating': False,
                'pos_h': None,    # left offset en EMU
                'pos_v': None,    # top offset en EMU
                'align_h': None,  # left/center/right
                'align_v': None,  # top/center/bottom
                'z_index': 0,
            }
            
            root = etree.fromstring(drawing_xml)
            
            # Buscar rId de la imagen
            blip_elem = root.find('.//a:blip', nsmap)
            if blip_elem is not None:
                rId = blip_elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rId:
                    info['rId'] = rId
            
            # Buscar extent (tamaño)
            extent = root.find('.//wp:extent', nsmap)
            if extent is not None:
                info['width_emu'] = int(extent.get('cx', 0))
                info['height_emu'] = int(extent.get('cy', 0))
            
            # Buscar posición horizontal
            pos_h = root.find('.//wp:positionH', nsmap)
            if pos_h is not None:
                info['floating'] = True
                align_h = pos_h.find('wp:align', nsmap)
                if align_h is not None:
                    info['align_h'] = align_h.text  # left, center, right
                else:
                    pos_offset = pos_h.find('wp:posOffset', nsmap)
                    if pos_offset is not None:
                        info['pos_h'] = int(pos_offset.text)
            
            # Buscar posición vertical
            pos_v = root.find('.//wp:positionV', nsmap)
            if pos_v is not None:
                info['floating'] = True
                align_v = pos_v.find('wp:align', nsmap)
                if align_v is not None:
                    info['align_v'] = align_v.text  # top, center, bottom
                else:
                    pos_offset = pos_v.find('wp:posOffset', nsmap)
                    if pos_offset is not None:
                        info['pos_v'] = int(pos_offset.text)
            
            # Buscar behindDoc
            behind = root.find('.//wp:behindDoc', nsmap)
            if behind is not None and behind.get('val') != '0':
                info['behind_text'] = True
            
            # Buscar z-index (relativeHeight)
            rel_h = root.find('.//wp:relativeHeight', nsmap)
            if rel_h is not None:
                info['z_index'] = int(rel_h.text or '0')
            
            return info
        
        def emu_to_pt(emu):
            """Convierte EMU (English Metric Units) a puntos."""
            return emu / 12700
        
        def emu_to_px(emu):
            """Convierte EMU a píxeles (aprox, 1pt = 1.333px)."""
            return int(emu / 12700 * 1.333)
        
        # ── Construir CSS de imágenes ─────────────────────────────────────
        images_css = []
        images_html_parts = []
        processed_images = set()
        
        # Procesar imágenes de los párrafos (XML completo del documento)
        doc_xml = etree.tostring(doc.element, encoding='unicode')
        doc_root = etree.fromstring(doc_xml.encode('utf-8'))
        
        # Buscar todos los drawings en el documento
        drawings = doc_root.findall('.//w:drawing', nsmap)
        # También buscar en inline shapes
        for para in doc.paragraphs:
            for run in para.runs:
                if run._element.findall('.//w:drawing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    for drawing in run._element.findall('.//w:drawing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        drawing_xml = etree.tostring(drawing, encoding='unicode')
                        img_info = extract_image_info(drawing_xml)
                        
                        rId = img_info['rId']
                        if not rId or rId not in image_rels:
                            continue
                        
                        img_data = image_rels[rId]
                        b64 = base64.b64encode(img_data['blob']).decode('utf-8')
                        img_id = f"img_{rId}_{len(processed_images)}"
                        processed_images.add(rId)
                        
                        # Determinar estilos
                        styles = []
                        styles.append(f"max-width:none")  # No limitar por CSS
                        
                        # Tamaño en puntos
                        w_pt = emu_to_pt(img_info['width_emu'])
                        h_pt = emu_to_pt(img_info['height_emu'])
                        
                        if img_info['floating']:
                            styles.append("position:absolute")
                            styles.append(f"width:{w_pt:.1f}pt")
                            styles.append(f"height:{h_pt:.1f}pt")
                            
                            # Posición horizontal
                            if img_info['align_h']:
                                if img_info['align_h'] == 'center':
                                    styles.append("left:50%")
                                    styles.append("transform:translateX(-50%)")
                                elif img_info['align_h'] == 'right':
                                    styles.append("right:0")
                                else:
                                    styles.append("left:0")
                            elif img_info['pos_h'] is not None:
                                styles.append(f"left:{emu_to_pt(img_info['pos_h']):.1f}pt")
                            
                            # Posición vertical
                            if img_info['align_v']:
                                if img_info['align_v'] == 'center':
                                    styles.append("top:50%")
                                elif img_info['align_v'] == 'bottom':
                                    styles.append("bottom:0")
                                else:
                                    styles.append("top:0")
                            elif img_info['pos_v'] is not None:
                                styles.append(f"top:{emu_to_pt(img_info['pos_v']):.1f}pt")
                            
                            # z-index
                            if img_info['behind_text']:
                                styles.append("z-index:-1")
                                styles.append("opacity:0.3")
                            else:
                                styles.append(f"z-index:{max(1, img_info['z_index'])}")
                        else:
                            # Imagen inline
                            styles.append(f"width:{w_pt:.1f}pt")
                            styles.append(f"height:{h_pt:.1f}pt")
                        
                        style_attr = "; ".join(styles)
                        
                        img_html = f'<img id="{img_id}" src="data:{img_data["mime"]};base64,{b64}" style="{style_attr}" />'
                        images_html_parts.append({
                            'rId': rId,
                            'html': img_html,
                            'floating': img_info['floating'],
                            'behind_text': img_info['behind_text'],
                        })
        
        # ── Construir HTML ────────────────────────────────────────────────
        html_parts = []
        html_parts.append('''
        <!DOCTYPE html>
        <html lang="es">
        <head>
        <meta charset="utf-8">
        <style>
          body {
            font-family: Calibri, 'Segoe UI', sans-serif;
            font-size: 11pt;
            max-width: 612pt;
            margin: 0 auto;
            padding: 56.7pt;
            color: #2d3748;
            background: #fff;
            position: relative;
            min-height: 100vh;
          }
          .preview-watermark {
            position: fixed;
            top: 12px;
            right: 16px;
            font-size: 10px;
            color: #a0aec0;
            z-index: 1000;
            font-family: 'Segoe UI', sans-serif;
            letter-spacing: .5px;
            background: #f7fafc;
            padding: 4px 10px;
            border-radius: 4px;
            border: 1px solid #e2e8f0;
          }
          .placeholder-highlight {
            background: #fefcbf;
            border: 1px solid #f6e05e;
            padding: 1px 4px;
            border-radius: 3px;
            font-weight: 600;
            color: #975a16;
            white-space: nowrap;
          }
          .overflow-warn {
            background: #fff5f5 !important;
            color: #e53e3e !important;
            border: 1px solid #feb2b2 !important;
            padding: 2px 4px !important;
            border-radius: 3px !important;
          }
          table { border-collapse: collapse; width: 100%; }
          td, th { padding: 4px 8px; border: 1px solid #cbd5e0; }
          img { max-width: 100%; height: auto; }
          p { margin: 0 0 6pt 0; line-height: 1.15; position: relative; }
          .text-right { text-align: right; }
          .text-center { text-align: center; }
          .text-justify { text-align: justify; }
          .page-container { position: relative; }
        </style>
        </head>
        <body>
        ''')
        
        if preview_on:
            html_parts.append('<div class="preview-watermark">👁️ VISTA PREVIA — Datos de ejemplo</div>')
        else:
            html_parts.append('<div class="preview-watermark">📍 PLANTILLA BASE — Placeholders visibles</div>')
        
        html_parts.append('<div class="page-container">')
        
        # ── Reemplazar placeholders en texto ──────────────────────────────
        def replace_in_text(text: str, highlight: bool = False) -> str:
            """Reemplaza placeholders {{campo}} por valores o los resalta."""
            if not text:
                return text
            
            result = text
            # Encontrar todos los placeholders
            pattern = r'\{\{(\w+)\}\}'
            
            if highlight and not preview_on:
                # Modo OFF: resaltar placeholders
                def highlight_match(m):
                    placeholder = m.group(0)
                    escaped = html_module.escape(placeholder)
                    return f'<span class="placeholder-highlight">{escaped}</span>'
                result = re.sub(pattern, highlight_match, result)
            elif preview_on and replacements:
                # Modo ON: reemplazar con valores
                for placeholder, valor in replacements.items():
                    escaped_valor = html_module.escape(str(valor))
                    # Detectar posible desborde
                    if len(str(valor)) > 40:
                        escaped_valor = f'<span class="overflow-warn" title="Posible desborde ({len(str(valor))} caracteres)">{escaped_valor}</span>'
                    result = result.replace(placeholder, escaped_valor)
            
            return result
        
        # ── Insertar imágenes flotantes detrás del texto al inicio ────────
        behind_images = [img for img in images_html_parts if img.get('behind_text')]
        for img in behind_images:
            html_parts.append(img['html'])
        
        # ── Procesar párrafos ─────────────────────────────────────────────
        for para in doc.paragraphs:
            alignment = para.alignment
            align_class = ""
            if alignment == 2: align_class = 'text-center'
            elif alignment == 3: align_class = 'text-right'
            elif alignment == 4: align_class = 'text-justify'
            
            # Verificar si este párrafo contiene una imagen inline
            has_inline_image = False
            inline_img_html = ""
            for run in para.runs:
                drawings_in_run = run._element.findall('.//w:drawing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                for drawing in drawings_in_run:
                    drawing_xml = etree.tostring(drawing, encoding='unicode')
                    img_info = extract_image_info(drawing_xml)
                    rId = img_info['rId']
                    if rId and rId in image_rels and not img_info.get('floating'):
                        img_data = image_rels[rId]
                        b64 = base64.b64encode(img_data['blob']).decode('utf-8')
                        w_pt = emu_to_pt(img_info['width_emu']) if img_info['width_emu'] else 200
                        h_pt = emu_to_pt(img_info['height_emu']) if img_info['height_emu'] else 100
                        inline_img_html = f'<img src="data:{img_data["mime"]};base64,{b64}" style="width:{w_pt:.1f}pt;height:{h_pt:.1f}pt;display:inline-block;vertical-align:middle;" />'
                        has_inline_image = True
                        break
                if has_inline_image:
                    break
            
            para_html = ""
            
            for run in para.runs:
                text = replace_in_text(run.text, highlight=True)
                escaped_text = html_module.escape(text) if text else ""
                
                styles = []
                if run.bold: styles.append("font-weight:bold")
                if run.italic: styles.append("font-style:italic")
                if run.underline: styles.append("text-decoration:underline")
                if run.font.size:
                    try:
                        size_pt = run.font.size / 12700
                        styles.append(f"font-size:{size_pt:.0f}pt")
                    except: pass
                if run.font.name:
                    styles.append(f"font-family:'{run.font.name}',sans-serif")
                
                style_attr = ";".join(styles)
                
                if preview_on and run.text:
                    # Reemplazar placeholders con valores (sin highlight)
                    replaced = replace_in_text(run.text, highlight=False)
                    if replaced != run.text:
                        para_html += replaced
                    elif escaped_text.strip():
                        para_html += f'<span style="{style_attr}">{escaped_text}</span>'
                    else:
                        para_html += escaped_text
                elif not preview_on and run.text:
                    # Modo OFF: mostrar con placeholders resaltados
                    highlighted = replace_in_text(run.text, highlight=True)
                    if highlighted != run.text:
                        para_html += highlighted
                    elif escaped_text.strip():
                        para_html += f'<span style="{style_attr}">{escaped_text}</span>'
                    else:
                        para_html += escaped_text
                else:
                    para_html += escaped_text
            
            # Si hay imagen inline, añadirla al párrafo
            if has_inline_image:
                # Determinar alineación de la imagen
                if align_class == 'text-center':
                    para_html = f'<div style="text-align:center">{inline_img_html}</div>' + para_html
                elif align_class == 'text-right':
                    para_html = f'<div style="text-align:right">{inline_img_html}</div>' + para_html
                else:
                    para_html = f'<div>{inline_img_html}</div>' + para_html
            
            is_list = para.style.name and ('List' in para.style.name or 'list' in para.style.name)
            
            if is_list:
                tag = "li"
                html_parts.append(f'<{tag} class="{align_class}">{para_html or "&nbsp;"}</{tag}>')
            else:
                style_name = para.style.name if para.style else ""
                if style_name and 'Heading' in style_name:
                    level = style_name.split()[-1]
                    try:
                        h_level = int(level)
                        tag = f"h{min(h_level, 6)}"
                    except: tag = "p"
                else:
                    tag = "p"
                html_parts.append(f'<{tag} class="{align_class}">{para_html or "&nbsp;"}</{tag}>')
        
        # ── Procesar tablas ───────────────────────────────────────────────
        for table in doc.tables:
            html_parts.append('<table>')
            for i, row in enumerate(table.rows):
                html_parts.append('<tr>')
                for cell in row.cells:
                    cell_html = []
                    for para in cell.paragraphs:
                        para_text = ""
                        for run in para.runs:
                            text = replace_in_text(run.text, highlight=True)
                            styles = []
                            if run.bold: styles.append("font-weight:bold")
                            if run.italic: styles.append("font-style:italic")
                            style_attr = ";".join(styles)
                            
                            if preview_on:
                                replaced = replace_in_text(run.text, highlight=False)
                                para_text += replaced if replaced != run.text else (f'<span style="{style_attr}">{html_module.escape(text) if text else ""}</span>')
                            else:
                                highlighted = replace_in_text(run.text, highlight=True)
                                para_text += highlighted if highlighted != run.text else (f'<span style="{style_attr}">{html_module.escape(text) if text else ""}</span>')
                        
                        align_class = ""
                        if para.alignment == 2: align_class = 'text-center'
                        elif para.alignment == 3: align_class = 'text-right'
                        cell_html.append(f'<p class="{align_class}" style="margin:0">{para_text or "&nbsp;"}</p>')
                    
                    tag = "th" if i == 0 else "td"
                    html_parts.append(f'<{tag}>{"".join(cell_html)}</{tag}>')
                html_parts.append('</tr>')
            html_parts.append('</table>')
        
        # ── Insertar imágenes flotantes delante del texto ─────────────────
        front_images = [img for img in images_html_parts if img.get('floating') and not img.get('behind_text')]
        for img in front_images:
            html_parts.append(img['html'])
        
        html_parts.append('</div>')  # page-container
        html_parts.append('</body></html>')
        
        full_html = "\n".join(html_parts)
        
        return {
            "success": True,
            "html": full_html,
            "preview_on": preview_on,
        }
    
    except ImportError as e:
        raise HTTPException(
            status_code=500,
            detail=f"Librería requerida no instalada: {str(e)}. Ejecuta: pip install python-docx lxml"
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generando vista previa: {str(e)}")

@router.post("/{plantilla_id}/preview-pdf")
def preview_plantilla_pdf(
    plantilla_id: int,
    body: PreviewPDFRequest,
    current_user: Usuario = Depends(get_current_active_user),
    db: Session = Depends(get_global_db),
):
    """
    Vista previa PDF usando el MAPEO real de plantilla_campos.
    - OFF: muestra {{campo_bd}} (placeholders del sistema)
    - ON:  muestra datos de ejemplo del sistema
    """
    import tempfile, shutil, subprocess, base64, re, zipfile
    from pathlib import Path
    from lxml import etree
    
    p = _get_plantilla_or_404(db, plantilla_id)
    
    docx_path = p.ruta_archivo
    if not docx_path or not _os.path.exists(docx_path):
        raise HTTPException(status_code=404, detail="Archivo .docx no encontrado")
    
    # ═══════════════════════════════════════════════════════════════
    # CARGAR MAPEO REAL DE LA BASE DE DATOS
    # ═══════════════════════════════════════════════════════════════
    campos_mapeo = db.query(PlantillaCampo).filter(
        PlantillaCampo.id_plantilla == plantilla_id
    ).all()
    
    # Mapeo: placeholder ({{campo}}) → campo_bd
    # Pero necesitamos el inverso: encontrar el MERGEFIELD original → placeholder
    # Los placeholders en plantilla_campos se guardan como "{{campo_bd}}"
    # y el campo_bd es el nombre real de la columna
    
    # Construir: campo_bd → placeholder completo
    bd_to_placeholder = {}
    bd_to_nombre = {}
    for c in campos_mapeo:
        # c.placeholder = "{{campo_bd}}" o "{{MERGEFIELD_ORIGINAL}}"
        # c.campo_bd = nombre real de la columna en BD
        bd = c.campo_bd
        bd_to_placeholder[bd] = f"{{{{{bd}}}}}"  # placeholder del sistema
        bd_to_nombre[bd] = bd
    
    print(f"[Preview] Mapeo cargado: {len(campos_mapeo)} campos")
    for c in campos_mapeo[:5]:
        print(f"  placeholder={c.placeholder} → campo_bd={c.campo_bd}")
    
    # ═══════════════════════════════════════════════════════════════
    # CREAR MAPEO: MERGEFIELD original → campo_bd del sistema
    # ═══════════════════════════════════════════════════════════════
    # El placeholder puede venir como "{{nombre}}" o como "MERGEFIELD nombre"
    # Necesitamos extraer el nombre limpio para hacer el match
    
    mergefield_to_bd = {}
    
    for c in campos_mapeo:
        placeholder = c.placeholder  # ej: "{{nombre}}" o "{{MERGEFIELD nombre}}"
        campo_bd = c.campo_bd
        
        # Extraer nombre limpio del placeholder
        limpio = placeholder.strip('{}').strip()
        # Si es "MERGEFIELD nombre", extraer "nombre"
        if limpio.upper().startswith('MERGEFIELD'):
            limpio = limpio.split(' ', 1)[1].strip() if ' ' in limpio else limpio
        
        mergefield_to_bd[limpio] = campo_bd
    
    print(f"[Preview] MERGEFIELD → BD: {mergefield_to_bd}")
    
    if not mergefield_to_bd:
        raise HTTPException(
            status_code=400,
            detail="Esta plantilla no tiene mapeo de campos. Mapea los placeholders primero."
        )
    
    temp_dir = tempfile.mkdtemp(prefix="trinnova_preview_")
    
    try:
        temp_docx = _os.path.join(temp_dir, "modified.docx")
        
        with zipfile.ZipFile(docx_path, 'r') as zin:
            with zipfile.ZipFile(temp_docx, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    
                    if item.filename.startswith('word/') and item.filename.endswith('.xml'):
                        try:
                            xml_str = data.decode('utf-8', errors='replace')
                            xml_str = _procesar_xml_con_mapeo(
                                xml_str,
                                body.preview_on,
                                body.placeholders,
                                mergefield_to_bd
                            )
                            data = xml_str.encode('utf-8')
                        except Exception as e:
                            print(f"[Preview] Error en {item.filename}: {e}")
                    
                    zout.writestr(item, data)
        
        # Convertir a PDF
        sistema = platform.system()
        if sistema == "Windows":
            posibles = [
                "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
                "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
            ]
            lo_path = "libreoffice"
            for ruta in posibles:
                if _os.path.exists(ruta):
                    lo_path = ruta
                    break
        else:
            lo_path = "libreoffice"
        
        comando = [
            lo_path, "--headless", "--norestore",
            "--convert-to", "pdf",
            "--outdir", temp_dir,
            temp_docx
        ]
        
        resultado = subprocess.run(comando, capture_output=True, text=True, timeout=60, shell=(sistema == "Windows"))
        
        if resultado.returncode != 0:
            raise HTTPException(status_code=500, detail=f"LibreOffice error: {resultado.stderr[:300]}")
        
        pdf_files = list(Path(temp_dir).glob("*.pdf"))
        if not pdf_files:
            raise HTTPException(status_code=500, detail="No se generó el PDF")
        
        with open(str(pdf_files[0]), 'rb') as f:
            pdf_bytes = f.read()
        
        pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
        
        registrar_log(
            db, current_user.id, "preview_plantilla_pdf",
            f"Preview {'ON' if body.preview_on else 'OFF'} plantilla {plantilla_id}",
            p.id_proyecto
        )
        
        return {
            "success": True,
            "pdf_base64": pdf_base64,
            "preview_on": body.preview_on,
            "plantilla_id": plantilla_id,
        }
    
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def _procesar_xml_con_mapeo(xml_str: str, preview_on: bool, placeholders: dict, mergefield_to_bd: dict) -> str:
    """
    Reemplaza el texto visible de los MERGEFIELD usando el mapeo.
    
    mergefield_to_bd = {'Clave_APA_1': 'clave_APA', 'DOMICILIO': 'calle', ...}
    placeholders = {'clave_APA': 'A-56789', 'calle': 'Calle Hidalgo 456', ...}
    
    - OFF: {{campo_bd}}  →  "{{clave_APA}}"
    - ON:  valor real    →  "A-56789"
    """
    from lxml import etree
    
    try:
        root = etree.fromstring(xml_str.encode('utf-8'))
    except:
        return xml_str
    
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    }
    
    instr_texts = root.findall('.//w:instrText', nsmap)
    
    for instr in instr_texts:
        text = instr.text.strip() if instr.text else ""
        
        if not text.upper().startswith('MERGEFIELD'):
            continue
        
        parts = text.split()
        if len(parts) < 2:
            continue
        
        field_name = parts[1].strip('"')
        
        # Buscar el campo_bd correspondiente en el mapeo
        campo_bd = mergefield_to_bd.get(field_name)
        
        if not campo_bd:
            # Si no hay mapeo, dejar como está
            continue
        
        # Buscar el run que contiene el texto visible
        instr_run = instr.getparent()
        if instr_run is None:
            continue
        
        parent = instr_run.getparent()
        if parent is None:
            continue
        
        all_runs = parent.findall('.//w:r', nsmap)
        
        separate_found = False
        value_run = None
        
        for run in all_runs:
            fld_chars = run.findall('.//w:fldChar', nsmap)
            for fc in fld_chars:
                fld_type = fc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType')
                if fld_type == 'separate':
                    separate_found = True
                    break
            
            if separate_found:
                t_elem = run.find('.//w:t', nsmap)
                if t_elem is not None and t_elem.text:
                    value_run = run
                    break
                
                fld_chars2 = run.findall('.//w:fldChar', nsmap)
                for fc in fld_chars2:
                    fld_type = fc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType')
                    if fld_type == 'end':
                        separate_found = False
                        break
        
        if value_run is None:
            continue
        
        t_elem = value_run.find('.//w:t', nsmap)
        if t_elem is None:
            continue
        
        # ═══ REEMPLAZAR USANDO EL MAPEO ═══
        if preview_on:
            # ON: mostrar valor de ejemplo
            valor = placeholders.get(campo_bd, f'[{campo_bd}]')
            t_elem.text = str(valor)
        else:
            # OFF: mostrar placeholder del sistema
            t_elem.text = f'{{{{{campo_bd}}}}}'
        
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    
    return etree.tostring(root, encoding='unicode')